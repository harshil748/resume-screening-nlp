"""Utilities for extracting text from uploaded resume files."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import BinaryIO


def _read_bytes(file_obj: BinaryIO) -> bytes:
    """Read bytes from a file-like object without assuming its current position."""
    if file_obj is None or not hasattr(file_obj, "read"):
        return b""

    try:
        if hasattr(file_obj, "seek"):
            file_obj.seek(0)
        data = file_obj.read()
        if isinstance(data, str):
            return data.encode("utf-8")
        return data or b""
    except Exception:
        return b""


def _extract_pdf_text(data: bytes) -> str:
    text_parts: list[str] = []

    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(BytesIO(data))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)
    except Exception:
        text_parts = []

    text = "\n".join(text_parts).strip()
    if text:
        return text

    try:
        import pdfplumber

        with pdfplumber.open(BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
    except Exception:
        return ""

    return "\n".join(text_parts).strip()


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(data))
        text_parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        for table in document.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts).strip()
    except Exception:
        return ""


def _extract_txt_text(data: bytes) -> str:
    try:
        return data.decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""


def extract_text(file_obj, filename: str) -> str:
    """Extract readable text from a PDF, DOCX, or TXT resume upload.

    The file type is detected from the filename extension. PDF extraction first
    uses PyPDF2 and falls back to pdfplumber when PyPDF2 produces no text. DOCX
    extraction includes both paragraphs and table cells. Any unsupported file
    type, unreadable input, or parser failure returns an empty string.
    """
    try:
        if not filename:
            return ""

        extension = Path(filename).suffix.lower()
        data = _read_bytes(file_obj)
        if not data:
            return ""

        if extension == ".pdf":
            return _extract_pdf_text(data)
        if extension == ".docx":
            return _extract_docx_text(data)
        if extension == ".txt":
            return _extract_txt_text(data)

        return ""
    except Exception:
        return ""
